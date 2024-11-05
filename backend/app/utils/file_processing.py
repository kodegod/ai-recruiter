import os
import uuid
import logging
from fastapi import UploadFile, HTTPException
import PyPDF2
import docx
from typing import Dict, Any
import json
import re

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}

def extract_text_from_file(file_path: str) -> Dict[str, Any]:
    """
    Extract text and metadata from various file types
    """
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return extract_text_from_txt(file_path)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_extension}"
            )
            
    except Exception as e:
        logger.error(f"Error extracting text from file: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error processing file: {str(e)}"
        )

def save_uploaded_file(file: UploadFile, directory: str = "uploads") -> str:
    """
    Save an uploaded file and return its path
    """
    try:
        # Create directory if it doesn't exist
        os.makedirs(directory, exist_ok=True)
        
        # Generate unique filename
        file_extension = os.path.splitext(file.filename)[1]
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        file_path = os.path.join(directory, unique_filename)
        
        # Save file
        file_content = file.file.read()
        with open(file_path, "wb") as buffer:
            buffer.write(file_content)
        
        # Reset file pointer
        file.file.seek(0)
        
        logger.info(f"File saved successfully: {file_path}")
        return file_path
        
    except Exception as e:
        logger.error(f"Error saving file: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error saving file: {str(e)}"
        )

async def process_jd_file(file: UploadFile) -> Dict[str, Any]:
    """
    Process job description file and extract structured information
    """
    try:
        # Check file extension
        file_extension = os.path.splitext(file.filename)[1].lower()
        if file_extension not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid file type. Allowed types are: {', '.join(ALLOWED_EXTENSIONS)}"
            )

        # Save file temporarily
        file_path = save_uploaded_file(file, directory="jd_uploads")
        
        try:
            # Extract text and metadata
            content = extract_text_from_file(file_path)
            
            # Extract job details
            job_details = extract_job_details(content["content"])
            
            return {
                "content": content["content"],
                "title": job_details.get("title", "Untitled Position"),
                "company": job_details.get("company", "Unknown Company"),
                "metadata": content.get("metadata", {})
            }
            
        finally:
            # Clean up temporary file
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Temporary file removed: {file_path}")
                
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing JD file: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error processing job description: {str(e)}"
        )

async def process_resume_file(file: UploadFile) -> Dict[str, Any]:
    """
    Process resume file and extract structured information
    """
    try:
        # Check file extension
        file_extension = os.path.splitext(file.filename)[1].lower()
        if file_extension not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Invalid file type. Allowed types are: {', '.join(ALLOWED_EXTENSIONS)}"
            )

        # Save file temporarily
        file_path = save_uploaded_file(file, directory="resume_uploads")
        
        try:
            # Extract text and metadata
            content = extract_text_from_file(file_path)
            
            # Extract candidate information with defaults
            candidate_info = extract_candidate_info(content["content"])
            
            processed_data = {
                "content": content["content"],
                "candidate_name": candidate_info.get("name") or "Unknown Candidate",
                "email": candidate_info.get("email", ""),  # Allow empty email, will be handled in upload endpoint
                "metadata": content.get("metadata", {}),
                "skills": [],  # Add empty skills list
                "experience": []  # Add empty experience list
            }
            
            logger.info(f"Processed resume data: {json.dumps(processed_data, default=str)[:200]}...")  # Log first 200 chars
            
            return processed_data
            
        finally:
            # Clean up temporary file
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Temporary file removed: {file_path}")
                
    except HTTPException as he:
        logger.error(f"HTTP Exception in process_resume_file: {str(he)}")
        raise he
    except Exception as e:
        logger.error(f"Error processing resume file: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error processing resume: {str(e)}"
        )

def extract_text_from_pdf(file_path: str) -> Dict[str, Any]:
    """
    Extract text and metadata from PDF files
    """
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            metadata = {}
            
            # Extract text from each page
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            # Extract metadata if available
            if reader.metadata:
                metadata = {
                    "title": reader.metadata.get("/Title", ""),
                    "author": reader.metadata.get("/Author", ""),
                    "created": reader.metadata.get("/CreationDate", "")
                }
            
            return {
                "content": text.strip(),
                "metadata": metadata
            }
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error processing PDF file: {str(e)}"
        )

def extract_text_from_docx(file_path: str) -> Dict[str, Any]:
    """
    Extract text and metadata from DOCX files
    """
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        # Extract metadata
        metadata = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "created": str(doc.core_properties.created) if doc.core_properties.created else ""
        }
        
        return {
            "content": text.strip(),
            "metadata": metadata
        }
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error processing DOCX file: {str(e)}"
        )

def extract_text_from_txt(file_path: str) -> Dict[str, Any]:
    """
    Extract text from TXT files
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
            return {
                "content": text.strip(),
                "metadata": {}
            }
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error processing TXT file: {str(e)}"
        )

def extract_job_details(content: str) -> Dict[str, str]:
    """
    Extract job title and company from content
    """
    # Common patterns for job titles and company names
    title_patterns = [
        r"Job Title:\s*(.*?)(?:\n|$)",
        r"Position:\s*(.*?)(?:\n|$)",
        r"Role:\s*(.*?)(?:\n|$)",
        r"^([A-Z][A-Za-z\s]+(?:Developer|Engineer|Manager|Analyst|Designer))(?:\n|$)"
    ]
    
    company_patterns = [
        r"Company:\s*(.*?)(?:\n|$)",
        r"Organization:\s*(.*?)(?:\n|$)",
        r"Employer:\s*(.*?)(?:\n|$)",
        r"@\s*([A-Za-z0-9\s]+)(?:\n|$)"
    ]
    
    title = ""
    company = ""
    
    # Find job title
    for pattern in title_patterns:
        match = re.search(pattern, content, re.MULTILINE | re.IGNORECASE)
        if match:
            title = match.group(1).strip()
            break
    
    # Find company name
    for pattern in company_patterns:
        match = re.search(pattern, content, re.MULTILINE | re.IGNORECASE)
        if match:
            company = match.group(1).strip()
            break
    
    return {
        "title": title,
        "company": company
    }

def extract_candidate_info(content: str) -> Dict[str, str]:
    """
    Extract candidate name and email from resume content
    """
    # Patterns for name and email
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    name_pattern = r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)'
    
    email = ""
    name = ""
    
    # Find email
    email_match = re.search(email_pattern, content)
    if email_match:
        email = email_match.group(0)
    
    # Find name (usually at the start of resume)
    lines = content.split('\n')
    for line in lines[:3]:  # Check first 3 lines
        name_match = re.match(name_pattern, line.strip())
        if name_match:
            name = name_match.group(0)
            break
    
    return {
        "name": name,
        "email": email
    }